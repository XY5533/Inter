"""
知识库文档解析模块
自动读取 knowledge_base/ 下的 PDF / Word / TXT 文件
"""

import os
from typing import List
from langchain_core.documents import Document

import config


def load_knowledge(kb_dir: str = None) -> List[Document]:
    """加载知识库目录中的所有文档，返回 Document 列表"""
    if kb_dir is None:
        kb_dir = config.KNOWLEDGE_BASE_DIR

    if not os.path.exists(kb_dir):
        os.makedirs(kb_dir, exist_ok=True)
        return []

    documents = []
    for filename in os.listdir(kb_dir):
        filepath = os.path.join(kb_dir, filename)
        if not os.path.isfile(filepath):
            continue

        ext = os.path.splitext(filename)[1].lower()
        try:
            if ext == ".txt":
                docs = _load_txt(filepath)
            elif ext == ".pdf":
                docs = _load_pdf(filepath)
            elif ext == ".docx":
                docs = _load_docx(filepath)
            else:
                print(f"  - {filename}: 不支持的格式")
                continue

            for doc in docs:
                doc.metadata["source"] = filename
            documents.extend(docs)
            print(f"  [OK] {filename} ({len(docs)} 个片段)")
        except Exception as e:
            print(f"  [FAIL] {filename}: {e}")

    return documents


def _load_txt(filepath: str) -> List[Document]:
    """读取 TXT 文件"""
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text)]


def _load_pdf(filepath: str) -> List[Document]:
    """读取 PDF 文件"""
    from pypdf import PdfReader
    reader = PdfReader(filepath)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text.strip():
            pages.append(Document(page_content=text.strip()))
    return pages


def _load_docx(filepath: str) -> List[Document]:
    """读取 Word 文件"""
    from docx import Document as DocxDocument
    doc = DocxDocument(filepath)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text)]
